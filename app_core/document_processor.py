"""
Document Processing Module
Handles text extraction from various file formats
"""
import logging
import re
from pathlib import Path
import fitz  # PyMuPDF
from docx import Document as DocxDocument
from openpyxl import load_workbook
from django.conf import settings

logger = logging.getLogger(__name__)


class DocumentProcessor:
    """
    Extract and clean text from various document formats
    """
    
    @staticmethod
    def extract_text(file_path, file_type):
        """
        Extract text from document based on file type
        
        Args:
            file_path (str): Path to the document
            file_type (str): Type of file (pdf, docx, txt, xlsx)
        
        Returns:
            str: Extracted text content
        """
        extractors = {
            'pdf': DocumentProcessor._extract_from_pdf,
            'docx': DocumentProcessor._extract_from_docx,
            'txt': DocumentProcessor._extract_from_txt,
            'xlsx': DocumentProcessor._extract_from_xlsx,
        }
        
        extractor = extractors.get(file_type.lower())
        if not extractor:
            raise ValueError(f"Unsupported file type: {file_type}")
        
        try:
            logger.info(f"Extracting text from {file_type} file: {file_path}")
            text = extractor(file_path)
            
            # Clean and normalize text
            text = DocumentProcessor._clean_text(text)
            
            logger.info(f"Successfully extracted {len(text)} characters")
            return text
            
        except Exception as e:
            logger.error(f"Error extracting text from {file_path}: {str(e)}")
            raise
    
    @staticmethod
    def _extract_from_pdf(file_path):
        """Extract text from PDF file using PyMuPDF"""
        text = []
        
        with fitz.open(file_path) as doc:
            for page_num, page in enumerate(doc, 1):
                page_text = page.get_text()
                if page_text.strip():
                    text.append(f"[Page {page_num}]\n{page_text}")
        
        return "\n\n".join(text)
    
    @staticmethod
    def _extract_from_docx(file_path):
        """Extract text from DOCX file"""
        doc = DocxDocument(file_path)
        
        text = []
        
        # Extract from paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                text.append(para.text)
        
        # Extract from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells]
                text.append(" | ".join(row_text))
        
        return "\n\n".join(text)
    
    @staticmethod
    def _extract_from_txt(file_path):
        """Extract text from plain text file"""
        encodings = ['utf-8', 'latin-1', 'cp1252']
        
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as f:
                    return f.read()
            except UnicodeDecodeError:
                continue
        
        # If all encodings fail, read as binary and decode with errors ignored
        with open(file_path, 'rb') as f:
            return f.read().decode('utf-8', errors='ignore')
    
    @staticmethod
    def _extract_from_xlsx(file_path):
        """Extract text from Excel file"""
        wb = load_workbook(file_path, read_only=True, data_only=True)
        
        text = []
        
        for sheet_name in wb.sheetnames:
            sheet = wb[sheet_name]
            text.append(f"[Sheet: {sheet_name}]")
            
            for row in sheet.iter_rows(values_only=True):
                row_text = [str(cell) if cell is not None else '' for cell in row]
                row_text = [cell for cell in row_text if cell.strip()]
                if row_text:
                    text.append(" | ".join(row_text))
        
        return "\n\n".join(text)
    
    @staticmethod
    def _clean_text(text):
        """
        Clean and normalize extracted text
        
        Args:
            text (str): Raw extracted text
        
        Returns:
            str: Cleaned text
        """
        if not text:
            return ""
        
        # Remove excessive whitespace
        text = re.sub(r'\s+', ' ', text)
        
        # Remove excessive newlines (keep max 2)
        text = re.sub(r'\n{3,}', '\n\n', text)
        
        # Remove control characters
        text = re.sub(r'[\x00-\x08\x0b-\x0c\x0e-\x1f\x7f-\x9f]', '', text)
        
        # Strip leading/trailing whitespace
        text = text.strip()
        
        return text
    
    @staticmethod
    def chunk_text(text, chunk_size=None, overlap=None):
        """
        Split text into overlapping chunks for RAG
        
        Args:
            text (str): Text to chunk
            chunk_size (int): Maximum characters per chunk
            overlap (int): Number of overlapping characters
        
        Returns:
            list: List of text chunks
        """
        # Ensure chunk_size and overlap are integers, with fallback to settings
        if chunk_size is None:
            chunk_size = settings.CHUNK_SIZE
        chunk_size = int(chunk_size) if not isinstance(chunk_size, int) else chunk_size
        
        if overlap is None:
            overlap = settings.CHUNK_OVERLAP
        overlap = int(overlap) if not isinstance(overlap, int) else overlap
        
        if not text:
            return []
            
        if len(text) <= chunk_size:
            return [text] if text else []
        
        chunks = []
        start = 0
        
        while start < len(text):
            # Find end of chunk
            end = start + chunk_size
            
            # If not at the end of text, try to break at sentence boundary
            if end < len(text):
                # Look for sentence endings within the last 20% of chunk
                search_start = end - int(chunk_size * 0.2)
                sentence_end = max(
                    text.rfind('. ', search_start, end),
                    text.rfind('! ', search_start, end),
                    text.rfind('? ', search_start, end),
                    text.rfind('\n', search_start, end)
                )
                
                if sentence_end != -1:
                    end = sentence_end + 1
            
            # Extract chunk
            chunk = text[start:end].strip()
            if chunk:
                chunks.append(chunk)
            
            # Move start position with overlap
            start = end - overlap
            
            # Ensure we make progress (avoid infinite loop)
            if chunks and start <= len(chunks[-1]):
                start = end
        
        logger.info(f"Split text into {len(chunks)} chunks (size={chunk_size}, overlap={overlap})")
        return chunks
    
    @staticmethod
    def get_text_preview(text, max_length=500):
        """
        Get a preview of text for display
        
        Args:
            text (str): Full text
            max_length (int): Maximum length of preview
        
        Returns:
            str: Text preview
        """
        # Ensure max_length is an integer
        max_length = int(max_length) if not isinstance(max_length, int) else max_length
        
        if not text:
            return ""
        
        if len(text) <= max_length:
            return text
        
        # Try to break at sentence
        preview = text[:max_length]
        last_sentence = max(
            preview.rfind('. '),
            preview.rfind('! '),
            preview.rfind('? ')
        )
        
        if last_sentence > max_length * 0.7:
            preview = preview[:last_sentence + 1]
        
        return preview + "..."